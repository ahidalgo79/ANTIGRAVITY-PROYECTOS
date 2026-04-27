#!/usr/bin/env python3
"""Revisor LaTeX+BibTeX para envío a Elsevier."""
import re, sys, time
from pathlib import Path
import bibtexparser
from bibtexparser.bparser import BibTexParser
from bibtexparser.customization import convert_to_unicode
from docx import Document

def extraer_citas(tex_path: str) -> set:
    content = Path(tex_path).read_text(encoding="utf-8", errors="ignore")
    matches = re.findall(r'\\cite[\w]*\*?{([^}]+)}', content)
    return {k.strip() for m in matches for k in m.split(',') if k.strip()}

def parsear_bib(bib_path: str) -> dict:
    parser = BibTexParser()
    parser.customization = convert_to_unicode
    with open(bib_path, encoding="utf-8") as f:
        db = bibtexparser.load(f, parser=parser)
    return {e["ID"]: e for e in db.entries}

def validar_campos_elsevier(entries: dict) -> list:
    req = {"article":["author","title","journal","year","volume","pages"],
           "inproceedings":["author","title","booktitle","year","pages"],
           "book":["author","title","publisher","year"]}
    problemas = []
    for eid, ent in entries.items():
        t = ent.get("ENTRYTYPE","misc").lower()
        faltan = [f for f in req.get(t, req.get("misc",[])) if f not in ent or not str(ent[f]).strip()]
        if faltan: problemas.append(f"⚠️ `{eid}` ({t}) falta: {', '.join(faltan)}")
        if "doi" not in ent and "url" not in ent and t in req:
            problemas.append(f"ℹ️ `{eid}` sin DOI/URL (recomendado)")
    return problemas

def generar_reporte(tex: str, bib: str) -> str:
    citas = extraer_citas(tex)
    bib_entries = parsear_bib(bib)
    claves = set(bib_entries.keys())
    indef = citas - claves
    huerf = claves - citas
    campos = validar_campos_elsevier(bib_entries)
    
    return f"""# 📄 Revisión para Elsevier
## 📊 Consistencia
- Citas en `.tex`: {len(citas)} | Entradas `.bib`: {len(bib_entries)}
- {'✅ Todas definidas' if not indef else f'❌ Indefinidas: `{", ".join(list(indef)[:5])}`'}
- {'✅ Sin huérfanos' if not huerf else f'⚠️ No citadas: {len(huerf)} entradas'}

## 📋 Campos Obligatorios
{chr(10).join(campos) if campos else '✅ Todos completos.'}

## 🚀 Checklist Pre-Envío
1. `bibtex main` → `xelatex main` (x2)
2. Corregir citas `❌` y añadir DOIs `ℹ️`
3. Usar `\\documentclass[review,preprint]{{elsarticle}}`
"""

def main(tex: str = "main.tex", bib: str = "refs.bib", fmt: str = "md"):
    print(f"🔍 Validando {Path(tex).name} + {Path(bib).name}...", flush=True)
    r = generar_reporte(tex, bib)
    print("\n" + "="*60 + "\n📄 REPORTE\n" + "="*60 + "\n" + r)
    
    out = Path("reports") / f"latex_check_{int(time.time())}"
    out.parent.mkdir(exist_ok=True)
    if fmt == "word":
        doc = Document()
        doc.add_heading("Revisión LaTeX/BibTeX", level=1)
        for l in r.splitlines():
            if l.startswith("# "): doc.add_heading(l[2:], level=2)
            elif l.startswith("## "): doc.add_heading(l[3:], level=3)
            elif l.startswith("- "): doc.add_paragraph(l[2:], style="List Bullet")
            elif l.strip(): doc.add_paragraph(l)
        out = out.with_suffix(".docx")
        doc.save(out)
    else:
        out = out.with_suffix(".md")
        out.write_text(r, encoding="utf-8")
    print(f"💾 Guardado: {out}")

if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument("--tex", default="main.tex")
    p.add_argument("--bib", default="references_clean.bib")
    p.add_argument("--format", choices=["md", "word"], default="md")
    args = p.parse_args()
    main(args.tex, args.bib, args.format)
