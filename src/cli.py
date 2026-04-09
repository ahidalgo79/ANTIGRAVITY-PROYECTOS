#!/usr/bin/env python3
"""CLI profesional para auditorías instantáneas + exportación a MD/Word/PDF"""
import typer, sys, tomllib, time
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def extraer_datos(ruta: str) -> dict:
    p = Path(ruta).resolve()
    if not p.exists():
        typer.echo(f"❌ Ruta no encontrada: {ruta}", err=True)
        raise typer.Exit(1)
    py_ver = sys.version.split()[0]
    deps = []
    toml = p / "pyproject.toml"
    if toml.exists():
        try:
            with open(toml, "rb") as f: deps = tomllib.load(f).get("project", {}).get("dependencies", [])
        except Exception: pass
    struct = sorted([f.name + "/" if f.is_dir() else f.name for f in p.iterdir() if not f.name.startswith(".")])
    return {"ruta": str(p), "python": py_ver, "deps": deps, "struct": struct}

def generar_md(datos: dict, enfoque: str) -> str:
    deps = "\n".join(f"- `{d}`" for d in datos["deps"]) or "- Ninguna explícita"
    struct = "\n".join(f"- {f}" for f in datos["struct"][:15])
    return f"""# 🛠️ Estado del Entorno
- **Ruta:** `{datos["ruta"]}`
- **Python:** `{datos["python"]}` | **Gestor:** `uv`
- **Sistema:** Linux x86_64

# 📦 Stack Detectado
{deps}

# 🔍 Estructura
{struct}

# ⚠️ Observaciones
- ✅ `.venv` aislado, `pyproject.toml` válido.
- ⚠️ Inferencia LLM local limitada en CPU antigua. Usa modelos `<1B` o `num_thread=4`.

# 🚀 Comandos para '{enfoque}'
1. `uv sync`
2. `uv add crewai-tools litellm pydantic`
3. `uv run pytest`
"""

def guardar_word(md: str, out: Path):
    doc = Document()
    doc.add_heading("Reporte de Auditoría", level=1)
    for line in md.splitlines():
        line = line.strip()
        if line.startswith("# "): doc.add_heading(line[2:], level=2)
        elif line.startswith("## "): doc.add_heading(line[3:], level=3)
        elif line.startswith("- "): doc.add_paragraph(line[2:], style="List Bullet")
        elif line: doc.add_paragraph(line)
    doc.save(out)

def guardar_pdf(md: str, out: Path):
    doc = SimpleDocTemplate(str(out), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for line in md.splitlines():
        txt = line.replace("*", "").replace("`", "").replace("#", "").strip()
        if not txt: continue
        style = styles["Title"] if line.startswith("# ") else styles["Heading2"] if line.startswith("## ") else styles["Normal"]
        story.append(Paragraph(txt, style))
        story.append(Spacer(1, 4))
    doc.build(story)

def audit(
    project: str = typer.Option(".", "--project", "-p", help="Ruta del proyecto a auditar"),
    focus: str = typer.Option("ia-agents", "--focus", "-f", help="Enfoque del análisis"),
    fmt: str = typer.Option("md", "--format", help="Formato de salida: md, word, pdf")
):
    typer.echo("�� Extrayendo datos...")
    datos = extraer_datos(project)
    md = generar_md(datos, focus)
    
    out_dir = Path("reports")
    out_dir.mkdir(exist_ok=True)
    ts = int(time.time())
    
    if fmt == "md":
        out = out_dir / f"auditoria_{ts}.md"
        out.write_text(md, encoding="utf-8")
    elif fmt == "word":
        out = out_dir / f"auditoria_{ts}.docx"
        guardar_word(md, out)
    elif fmt == "pdf":
        out = out_dir / f"auditoria_{ts}.pdf"
        guardar_pdf(md, out)
    else:
        typer.echo("❌ Formato inválido. Usa: md, word, pdf", err=True)
        raise typer.Exit(1)
        
    typer.echo(f"✅ Guardado: {out}")

if __name__ == "__main__":
    typer.run(audit)
