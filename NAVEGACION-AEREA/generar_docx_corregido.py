from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── PORTADA ───
for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD POLITÉCNICA DE CHIHUAHUA')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Navegación Aérea — Unidad 2')
run.font.size = Pt(14)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Planificación y Ejecución de Vuelo con Cartas Aeronáuticas y Garmin G1000')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Ruta: Mazatlán (MMMZ) – Ciudad Obregón (MMCN) – Ida y Vuelta')
run.font.size = Pt(13)

doc.add_paragraph('')

lines = [
    'Nombre completo y matrícula:',
    '1.',
    '2.',
    '3.',
    '4.',
    '',
    'Grupo:',
    'Fecha entrega:',
]
for l in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(l)
    run.font.size = Pt(12)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nombre del archivo a entregar: U2_Act2_GrupoX_EquipoY_MMMZ_MMCN.pdf')
run.bold = True
run.font.size = Pt(11)

doc.add_page_break()

# ─── 1. INTRODUCCIÓN ───
h = doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'Las cartas aeronáuticas son documentos esenciales para la planificación y ejecución segura '
    'de cualquier vuelo. En esta actividad, el equipo realizará un vuelo simulado de ida y vuelta '
    'entre Mazatlán (MMMZ) y Ciudad Obregón (MMCN) en Microsoft Flight Simulator X (FSX) a bordo '
    'de una aeronave equipada con el sistema aviónico Garmin G1000. El objetivo es identificar, '
    'interpretar y aplicar en tiempo real los diferentes tipos de cartas que intervienen en cada '
    'fase del vuelo, cruzando esta información con lo que muestra el equipo de navegación.'
)

# ─── 2. OBJETIVO ───
doc.add_heading('2. Objetivo de Aprendizaje', level=1)
doc.add_paragraph('Al finalizar la actividad, el alumno será capaz de:')
objs = [
    'Buscar, descargar e interpretar correctamente 5 tipos de cartas aeronáuticas reales (Aeródromo, SID, En-ruta, STAR y Aproximación) desde fuentes oficiales, específicamente para los aeropuertos MMMZ y MMCN.',
    'Programar un plan de vuelo completo (incluyendo procedimientos de salida y llegada) en el FMS del Garmin G1000 para la ruta MMMZ ↔ MMCN.',
    'Ejecutar un vuelo simulado respetando las altitudes, rumbos y restricciones publicadas en las cartas, documentando el proceso mediante bitácora, reporte y evidencia en video.',
]
for o in objs:
    doc.add_paragraph(o, style='List Number')

# ─── 3. MATERIALES ───
doc.add_heading('3. Materiales y Recursos Requeridos', level=1)
materiales = [
    'Simulador: Microsoft Flight Simulator X (FSX) o Prepar3D. Aeronave obligatoria: Cessna 172 SP con Garmin G1000 o Beechcraft Baron G58 con Garmin G1000.',
    'Cartas Aeronáuticas (Reales y vigentes): Deben descargarse de fuentes oficiales o confiables como: eAIP México (https://www.aip.gob.mx/) (Buscar cartas de MMMZ y MMCN), SkyVector (https://skyvector.com/) (Configurar región a México), AIRMATE.',
    'Software de captura: OBS Studio, Xbox Game Bar o NVIDIA ShadowPlay para grabar la pantalla del simulador.',
    'Procesador de texto: Word o Google Docs para el reporte escrito (exportar finalmente a PDF).',
    'Bitácora de vuelo: Debe imprimirse y llenarse a mano con bolígrafo durante o inmediatamente después del vuelo (ver Sección 6).',
]
for m in materiales:
    doc.add_paragraph(m, style='List Bullet')

# ─── 4. INSTRUCCIONES ───
doc.add_heading('4. Instrucciones Detalladas', level=1)

# Fase 1
doc.add_heading('Fase 1: Planificación en Tierra (Antes de encender el simulador)', level=2)
doc.add_paragraph(
    'El equipo debe trabajar en conjunto para definir la ruta. No se permite volar "Direct-To" '
    'en todo el trayecto; deben seguir airways o una ruta RNAV publicada.'
)

p = doc.add_paragraph()
run = p.add_run('1. Selección de Aeropuertos:')
run.bold = True
items = [
    'Vuelo de ida: MMMZ (Mazatlán) → MMCN (Ciudad Obregón)',
    'Vuelo de regreso: MMCN (Ciudad Obregón) → MMMZ (Mazatlán)',
]
for i in items:
    doc.add_paragraph(i, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('2. Análisis de Cartas:')
run.bold = True
doc.add_paragraph(
    'Descarguen e impriman las siguientes cartas para ambos aeropuertos. '
    'Para cada una, identifiquen y anoten los datos críticos:'
)

# Tabla de cartas
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['#', 'Tipo de Carta', 'Qué debe hacer el equipo']
cells_data = [
    ['1', 'Carta de Aeródromo (AD)', 'Identificar pista en uso, elevación del aeropuerto y ubicación del estacionamiento (Parking) de salida para MMMZ y MMCN.'],
    ['2', 'SID — Salida por Instrumentos', 'Elegir una salida para la pista activa de MMMZ. Identificar: nombre, primer waypoint, rumbo inicial y altitud de restricción.'],
    ['3', 'Carta En-ruta (Low / High)', 'Trazar la ruta entre MMMZ y MMCN. Identificar al menos 2 VORs/NDBs, 2 intersecciones (fixes) y el airway que conecta ambos aeropuertos.'],
    ['4', 'STAR — Llegada Estándar', 'Seleccionar una llegada a MMCN. Identificar: transición, último waypoint antes de la aproximación y altitud esperada.'],
    ['5', 'Carta de Aproximación', 'Elegir ILS, VOR o RNAV/GPS para la pista esperada en MMCN. Identificar: frecuencia del ayudante, curso final, altitud de decisión (DA) o altitud mínima de descenso.'],
]
for j, h in enumerate(headers):
    cell = table.cell(0, j)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
for i, row_data in enumerate(cells_data):
    for j, text in enumerate(row_data):
        table.cell(i+1, j).text = text

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('3. Plan de Vuelo:')
run.bold = True
doc.add_paragraph(
    'Anoten en la bitácora (Sección 6) la ruta completa en formato ICAO. '
    'Ejemplo (ruta ilustrativa, debe ser la elegida por el equipo): MMMZ DCT LMM UA301 MMCN'
)

# Fase 2
doc.add_heading('Fase 2: Ejecución del Vuelo Simulado en FSX con Garmin G1000', level=2)

doc.add_paragraph(
    '4. Configuración: Inicien la grabación de video antes de encender los motores. '
    'La grabación debe mostrar claramente las pantallas del G1000 (PFD y MFD).'
)

doc.add_paragraph(
    '5. Programación del FMS (G1000):'
)
fms_steps = [
    'a) Presionar el botón FPL.',
    'b) Ingresar aeropuerto de salida (MMMZ) y destino (MMCN).',
    'c) Presionar PROC para seleccionar e insertar la SID elegida para MMMZ.',
    'd) Insertar manualmente los waypoints de la ruta En-ruta (si no se cargan automáticamente).',
    'e) Presionar PROC nuevamente para pre-cargar la STAR y la Aproximación a MMCN. Activar con ACTV solo cuando el ATC o la fase del vuelo lo indiquen.',
]
for s in fms_steps:
    doc.add_paragraph(s)

doc.add_paragraph(
    '6. Ejecución y Capturas: Durante el vuelo, asegúrense de que el video capture '
    '(o tomen capturas de pantalla adicionales para el reporte) los siguientes momentos clave:'
)
capturas = [
    'Momento 1 (Salida): Pantalla MFD mostrando la SID activa y el avión siguiendo el rumbo.',
    'Momento 2 (Crucero): Pantalla PFD mostrando altitud y rumbo estabilizados, y MFD mostrando la posición sobre el Airway.',
    'Momento 3 (Transición a Llegada): Pantalla MFD mostrando la activación de la STAR a MMCN.',
    'Momento 4 (Aproximación Final): Pantalla PFD mostrando el localizador/glideslope (ILS) o la barra de desviación (RNAV) centrada.',
    'Momento 5 (Finalización): Avión detenido en el parking de MMCN, motores apagados.',
]
for c in capturas:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph(
    '7. Regreso: Repetir el proceso para el vuelo de regreso MMCN → MMMZ, '
    'utilizando una SID diferente desde MMCN, una ruta en-ruta distinta si es posible '
    'y una STAR / aproximación diferente a MMMZ.'
)

# Fase 3
doc.add_heading('Fase 3: Llenado de Bitácora y Elaboración del Reporte', level=2)
doc.add_paragraph(
    '8. Bitácora: Completen la bitácora impresa a mano con los datos reales obtenidos del simulador '
    '(no los planeados, sino los ejecutados). Debe estar firmada por el "Piloto al Mando" y el '
    '"Copiloto". Escanearla o tomarle una foto de alta calidad.'
)
doc.add_paragraph(
    '9. Reporte Escrito: Elaboren un documento en PDF siguiendo estrictamente la estructura de la Sección 5.'
)

# ─── 5. ESTRUCTURA DEL REPORTE ───
doc.add_heading('5. Estructura Obligatoria del Reporte Escrito (Máx. 5 MB)', level=1)
doc.add_paragraph('El reporte debe contener las siguientes secciones en este orden:')
secciones = [
    'Portada: Datos de la universidad, materia, actividad, integrantes, grupo y fecha.',
    'Resumen del Plan de Vuelo: Tabla con aeropuertos (MMMZ y MMCN), aeronave, ruta completa (string de navegación), altitud de crucero y aeropuerto alternativo seleccionado.',
    'Análisis de Cartas Aeronáuticas: Por cada una de las 5 cartas (para el vuelo de ida MMMZ→MMCN): insertar una imagen clara de la carta y texto breve explicando qué dato crítico extrajeron (ej. "De la carta de aproximación ILS RWY 04 de MMCN, tomamos la frecuencia 110.3 y la altitud de decisión de 2400 ft").',
    'Análisis del Garmin G1000: Insertar las 5 capturas de pantalla solicitadas en la Fase 2. Debajo de cada captura, explicar qué están mostrando los instrumentos y cómo coinciden con la carta aeronáutica.',
    'Conclusiones del Equipo: Mínimo 2 párrafos reflexionando sobre la importancia de la verificación cruzada (cross-check) entre la carta en papel/digital y lo que muestra el G1000, y los errores comunes que evitaron.',
    'Enlace al Video: Hipervínculo visible y funcional al video de YouTube (No listado) o Google Drive. El video debe tener una duración mínima de 5 minutos y, de preferencia, incluir audio del equipo narrando brevemente las transiciones en la ruta MMMZ-MMCN.',
    'Referencias: Formato APA 7ª edición (incluir la cita de las cartas del AIP de México o SkyVector para MMMZ y MMCN).',
]
for i, s in enumerate(secciones, 1):
    doc.add_paragraph(f'{i}. {s}')

# ─── 6. BITÁCORA ───
doc.add_heading('6. Bitácora de Vuelo (Para imprimir y llenar a mano)', level=1)
doc.add_paragraph(
    'Nota: Esta hoja debe anexarse escaneada o fotografiada al final del reporte PDF.'
)

# Tabla info general
t = doc.add_table(rows=2, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = 'Aeropuerto de Salida: Mazatlán'
t.cell(0,1).text = 'ICAO: MMMZ'
t.cell(0,2).text = 'Aeropuerto Destino: Ciudad Obregón'
t.cell(0,3).text = 'ICAO: MMCN'
t.cell(1,0).text = 'Aeronave (Modelo FSX):'
t.cell(1,1).text = 'Matrícula Sim:'
t.cell(1,2).text = 'Altitud de Crucero Planeada:'
t.cell(1,3).text = 'Aeropuerto Alternativo:'

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Registro de Tramos (Vuelo de Ida: MMMZ → MMCN)')
run.bold = True

t2 = doc.add_table(rows=2, cols=5)
t2.style = 'Table Grid'
h2 = ['Tramo / Waypoint', 'Carta de Ref.', 'Altitud (ft)', 'Rumbo (°)', 'Dist. (nm)']
for j, hh in enumerate(h2):
    run = t2.cell(0, j).paragraphs[0].add_run(hh)
    run.bold = True
for j in range(5):
    t2.cell(1, j).text = ''

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Registro de Tramos (Vuelo de Regreso: MMCN → MMMZ)')
run.bold = True

t3 = doc.add_table(rows=2, cols=4)
t3.style = 'Table Grid'
h3 = ['Tramo / Waypoint', 'Carta de Ref.', 'Altitud (ft)', 'Rumbo (°)']
for j, hh in enumerate(h3):
    run = t3.cell(0, j).paragraphs[0].add_run(hh)
    run.bold = True
for j in range(4):
    t3.cell(1, j).text = ''

doc.add_paragraph('')
doc.add_paragraph('Firma del Piloto al Mando:')
doc.add_paragraph('Firma del Copiloto:')
doc.add_paragraph('Hora Total de Bloque (Sim): _______ hrs')

# ─── 7. RÚBRICA ───
doc.add_heading('7. Criterios de Evaluación (Rúbrica)', level=1)

t4 = doc.add_table(rows=6, cols=3)
t4.style = 'Table Grid'
rh = ['Criterio', 'Descripción Detallada', 'Puntaje']
for j, hh in enumerate(rh):
    run = t4.cell(0, j).paragraphs[0].add_run(hh)
    run.bold = True
rubrica = [
    ['1. Bitácora de Vuelo', 'Está impresa, llenada a mano, firmada y los datos coinciden lógicamente con el reporte y el video. Incluye vuelo de ida (MMMZ→MMCN) y regreso (MMCN→MMMZ).', '20'],
    ['2. Uso de Cartas', 'Identifica y describe correctamente las 5 cartas para ambos aeropuertos. Las imágenes de las cartas son legibles y el texto explica el dato crítico extraído de cada una.', '20'],
    ['3. Interpretación G1000', 'Las capturas de pantalla del PFD/MFD están correctamente etiquetadas. El análisis explica claramente la relación entre lo que dice la carta y lo que muestra el instrumento en la ruta Mazatlán–Ciudad Obregón.', '20'],
    ['4. Reporte Escrito', 'Sigue la estructura solicitada, redacción profesional, sin faltas de ortografía, conclusiones reflexivas y referencias APA 7. Peso del archivo < 5 MB.', '20'],
    ['5. Evidencia en Video', 'El video (mín. 5 min) muestra claramente las pantallas del G1000 durante las transiciones clave. Se aprecia la programación del FMS y la ejecución de la aproximación en MMMZ y MMCN.', '20'],
]
for i, rd in enumerate(rubrica):
    for j, txt in enumerate(rd):
        t4.cell(i+1, j).text = txt
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('TOTAL')
run.bold = True
run = p.add_run('                                                          100')

# ─── 8. CHECKLIST ───
doc.add_heading('8. Puntos Críticos para Considerar (Checklist antes de entregar)', level=1)
checklist = [
    '¿El video es accesible? (Si es Drive, verificar permisos en "Cualquier persona con el enlace puede ver").',
    '¿Las capturas de pantalla del G1000 son legibles? (Si no, usar zoom o añadir anotaciones con flechas).',
    '¿Se respetó el límite de 5 MB del PDF? (Comprimir imágenes si es necesario).',
    '¿Se declaró el uso de herramientas de IA en el reporte, si se utilizaron para corrección de estilo? (Obligatorio por honestidad académica).',
    '¿La bitácora está firmada a mano?',
    '¿Todos los códigos ICAO en el reporte y bitácora son MMMZ y MMCN? (sin confusiones con MMCS o MMCU).',
]
for c in checklist:
    doc.add_paragraph(c, style='List Bullet')

# Guardar
output_path = '/home/andres/Documentos/ANTIGRAVITY-PROYECTOS/NAVEGACION-AEREA/Actividad_2_U2_Cartas_Aeronauticas.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
