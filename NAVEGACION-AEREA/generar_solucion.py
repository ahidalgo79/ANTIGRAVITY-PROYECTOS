from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_header_table(doc, title, subtitle=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1B3A6B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.color.rgb = RGBColor(0xE8, 0xF0, 0xFF)
        r2.font.size = Pt(11)
        r2.font.name = 'Arial'
    doc.add_paragraph()

# ========== TITLE ==========
add_header_table(doc, "SOLUCIÓN DETALLADA — EXAMEN UNIDAD 2", "Cartas Aeronáuticas • Radionavegación • Velocidades y Altimetría • Uso del E6B")

p = doc.add_paragraph()
run = p.add_run("Instrucciones para el docente:")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

items = [
    "Cada respuesta correcta tiene el puntaje indicado. Se sugiere otorgar puntaje parcial solo si el procedimiento está bien planteado aunque el resultado numérico sea incorrecto.",
    "En la Sección D, valorar el procedimiento paso a paso (50%) y el resultado final (50%).",
    "Puntaje total: 100 puntos. Mínimo aprobatorio: 70 puntos."
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.name = 'Arial'

doc.add_paragraph()

# ========== SECTION A ==========
add_header_table(doc, "SECCIÓN A: Cartas Aeronáuticas — Anexo 4 OACI", "20 puntos")

# P1
p = doc.add_paragraph()
run = p.add_run("P1. Relacionar columnas")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Carta", "Respuesta", "Descripción"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "D6E4F0")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'

data = [
    ("A) SID", "C", "Proporciona info. para la salida normalizada desde el despegue hasta la ruta."),
    ("B) STAR", "B", "Proporciona información para pasar de fase de ruta a la aproximación."),
    ("C) En-Ruta", "A", "Facilita la navegación a lo largo de rutas ATS (FIR)."),
    ("D) Plano de Aeródromo", "D", "Muestra las instalaciones terrestres, pistas y plataforma desde vista general."),
    ("E) IAC", "E", "Procedimiento de aproximación por instrumentos a la pista prevista."),
]
for i, (letter, ans, desc) in enumerate(data):
    row = table.rows[i + 1]
    for j, txt in enumerate([letter, f" ( {ans} )", desc]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

doc.add_paragraph()

# P2
p = doc.add_paragraph()
run = p.add_run("P2. Opción múltiple (10 pts — 2 pts c/u)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

mcq = [
    ("2.1", "Cartas obligatorias del Anexo 4 OACI",
     [
         "a) Tipo A, Carta Topográfica Precisión, Carta En Ruta, IAC, Plano Aeródromo y Mundial.",
         "b) Tipo B, Movimientos en tierra, Estacionamiento y Atraque.",
         "c) Carta de Área, SID, STAR, Aproximación Visual.",
     ],
     "a",
     "El Anexo 4 OACI (Cartas Aeronáuticas) especifica 6 cartas obligatorias: Tipo A (o Planeamiento), Carta Topográfica de Precisión, Carta En Ruta, Carta de Aproximación por Instrumentos (IAC), Plano de Aeródromo y Carta Mundial Aeronáutica (1:1,000,000)."),
    ("2.2", "Función del Plano de Estacionamiento y Atraque de Aeronaves",
     [
         "a) Movimiento general entre pistas y plataformas.",
         "b) Movimiento entre rodajes de acceso y puestos de estacionamiento específicos.",
         "c) Identificar obstáculos en el área de despegue.",
     ],
     "b",
     "El Plano de Estacionamiento y Atraque detalla los puestos de estacionamiento específicos y las rutas de acceso desde los rodajes, no el movimiento general entre pistas y plataformas."),
    ("2.3", "Tipo de fondo en zonas costeras seguras en cartas náuticas",
     ["a) Vegetación", "b) Arena", "c) Bosque"],
     "b",
     "En cartas náuticas aeronáuticas, las zonas costeras seguras se representan con fondo color arena/beige (tierra firme), mientras que el agua se representa en azul."),
    ("2.4", "Escala principal de las Cartas Seccionales Aeronáuticas",
     ["a) 1:500,000", "b) 1:1,000,000", "c) 1:100,000"],
     "a",
     "La escala estándar de las Cartas Seccionales Aeronáuticas (Sectional Charts) es 1:500,000, que ofrece un balance entre nivel de detalle y cobertura geográfica para navegación VFR."),
    ("2.5", "Método de navegación basado en cálculos matemáticos sin referencias visuales",
     ["a) Radionavegación", "b) Navegación visual", "c) Navegación a estima"],
     "c",
     "La navegación a estima (Dead Reckoning) se basa exclusivamente en cálculos de velocidad, tiempo, rumbo y deriva, sin depender de referencias visuales externas ni señales de radio."),
]

for qnum, question, options, answer, explanation in mcq:
    p = doc.add_paragraph()
    run = p.add_run(f"{qnum}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run(question)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    for opt in options:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(opt)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run(f"✅ Respuesta correcta: {answer}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

    p = doc.add_paragraph()
    run = p.add_run(f"Justificación: {explanation}")
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.italic = True
    doc.add_paragraph()

# ========== SECTION B ==========
add_header_table(doc, "SECCIÓN B: Sistemas y Radionavegación", "20 puntos")

p = doc.add_paragraph()
run = p.add_run("P3. Identificación de sistemas (20 pts — 4 pts c/u)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

systems = [
    ("3.1", "Sistema que proporciona información de distancia (slant range)",
     ["a) VOR", "b) DME", "c) ADF"],
     "b",
     "DME (Distance Measuring Equipment) mide la distancia oblicua (slant range) entre la aeronave y la estación mediante el tiempo de retorno de una señal de radio UHF."),
    ("3.2", "Sistema satelital para determinar posición tridimensional",
     ["a) GPS", "b) INS", "c) VOR"],
     "a",
     "GPS (Global Positioning System) utiliza una constelación de satélites para determinar coordenadas tridimensionales (latitud, longitud y altitud) mediante trilateración."),
    ("3.3", "Sistema para aproximación en baja visibilidad (localizador + senda de planeo)",
     ["a) ILS", "b) DME", "c) ADF"],
     "a",
     "ILS (Instrument Landing System) consta de dos componentes: el localizador (LOC) para guía lateral y la senda de planeo (GS) para guía vertical, permitiendo aproximaciones de precisión CAT I/II/III."),
    ("3.4", "Sistema VHF que permite volar radiales magnéticos desde/hacia la estación",
     ["a) VOR", "b) INS", "c) Satélite"],
     "a",
     "VOR (VHF Omnidirectional Range) emite señales en banda VHF (108-118 MHz) que permiten al piloto determinar y seguir radiales magnéticos desde o hacia la estación."),
    ("3.5", "Sistema de navegación inercial (giroscopios + acelerómetros, sin señales externas)",
     ["a) INS", "b) ILS", "c) VOR"],
     "a",
     "INS (Inertial Navigation System) es un sistema autónomo que utiliza giroscopios para mantener orientación y acelerómetros para medir desplazamientos, sin requerir señales externas de radio ni satélites."),
]

for qnum, question, options, answer, explanation in systems:
    p = doc.add_paragraph()
    run = p.add_run(f"{qnum}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run(question)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    for opt in options:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(opt)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run(f"✅ Respuesta correcta: {answer}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

    p = doc.add_paragraph()
    run = p.add_run(f"Justificación: {explanation}")
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.italic = True
    doc.add_paragraph()

# ========== SECTION C ==========
add_header_table(doc, "SECCIÓN C: Velocidades y Altimetría", "20 puntos")

p = doc.add_paragraph()
run = p.add_run("P4. Verdadero o Falso (20 pts — 2 pts c/u)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

tf_data = [
    ("1.", "IAS (Indicated Airspeed) es la velocidad que el piloto lee directamente en su instrumento.",
     "V", "Verdadero. IAS es la lectura directa del anemómetro sin correcciones, basada en la presión dinámica medida por el tubo Pitot."),
    ("2.", "TAS (True Airspeed) es la velocidad real de la aeronave con respecto a la masa de aire circundante.",
     "V", "Verdadero. TAS es la velocidad real relativa a la masa de aire, corregida por altitud y temperatura. Es la que se usa para navegación."),
    ("3.", "GS (Ground Speed) es igual a la TAS sumando o restando el efecto del viento.",
     "V", "Verdadero. GS = TAS ± componente de viento. Con viento de cola la GS aumenta; con viento en contra disminuye."),
    ("4.", "V1 es la velocidad de decisión durante la carrera de despegue.",
     "V", "Verdadero. V1 es la velocidad límite por debajo de la cual se puede abortar el despegue de forma segura; por encima de V1 se debe continuar."),
    ("5.", "QNH es la presión atmosférica ajustada para que el altímetro indique cero en el terreno.",
     "F", "Falso. QNH se ajusta para que el altímetro indique la altitud del aeródromo sobre el nivel del mar (elevación). QFE es el ajuste que hace que el altímetro indique cero en la pista."),
    ("6.", "Al ajustar el altímetro a QNE (1013.25 hPa) se lee la altitud de presión (niveles de vuelo).",
     "V", "Verdadero. QNE (1013.25 hPa / 29.92 inHg) es la referencia estándar por encima de la altitud de transición, y el altímetro indica altitud de presión, usada para Flight Levels."),
    ("7.", "La altitud de densidad disminuye cuando la temperatura exterior aumenta.",
     "F", "Falso. La altitud de densidad AUMENTA cuando la temperatura exterior aumenta. A mayor temperatura, el aire es menos denso, lo que se traduce en una mayor altitud de densidad."),
    ("8.", "La altura es la distancia vertical desde el nivel medio del mar.",
     "F", "Falso. La altura es la distancia vertical desde un punto de referencia en tierra (AGL — Above Ground Level). La altitud es la distancia vertical desde el nivel medio del mar (MSL)."),
    ("9.", "Vmcg es la velocidad mínima de control en tierra si falla un motor.",
     "V", "Verdadero. Vmcg (Minimum Control Ground) es la velocidad mínima durante la carrera de despegue en tierra en la que se puede controlar la aeronave con un motor crítico inoperativo."),
    ("10.", "La velocidad GS es la utilizada para estimar la hora de llegada (ETA).",
     "V", "Verdadero. La Ground Speed es la velocidad real sobre el terreno, por lo que combinada con la distancia permite calcular el tiempo estimado de llegada (ETA)."),
]

for number, statement, answer, explanation in tf_data:
    p = doc.add_paragraph()
    run = p.add_run(f"{number}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run(statement)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run(f"  ( {answer} )")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

    p = doc.add_paragraph()
    run = p.add_run(f"  {explanation}")
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.italic = True
    doc.add_paragraph()

# ========== SECTION D ==========
add_header_table(doc, "SECCIÓN D: Planificación con Computador de Vuelo E6B", "40 puntos")

p = doc.add_paragraph()
run = p.add_run("📌  CASO PRÁCTICO: BÚSQUEDA Y RESCATE EN ZONA MONTAÑOSA")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

p = doc.add_paragraph()
run = p.add_run("Datos de la misión:")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

data_items = [
    "Aeronave: Cessna 182T Skylane",
    "TAS de crucero de búsqueda: 120 kts",
    "Consumo de combustible: 10 gal/h",
    "Combustible utilizable a bordo: 56 gal (tanques llenos)",
    "Condiciones: VFR marginal, altitud de presión 8,500 ft, OAT = +25 °C",
    "Viento reportado: 220° / 25 kts",
    "Rumbo de búsqueda asignado (TC): 270°",
    "Duración de la búsqueda: 90 min",
    "Distancia de regreso a base: 54 NM (GS de regreso = 135 kts)",
    "Reserva legal requerida: 45 min",
]
for item in data_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.name = 'Arial'

doc.add_paragraph()

# P5
p = doc.add_paragraph()
run = p.add_run("P5. Cálculo de WCA y GS (12 pts)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

p = doc.add_paragraph()
run = p.add_run("Datos: Viento 220°/25 kts | TAS 120 kts | TC 270°")
run.font.size = Pt(11)
run.font.name = 'Arial'
run.italic = True

# Cálculo del WCA
p = doc.add_paragraph()
run = p.add_run("1. Cálculo del WCA (Wind Correction Angle):")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Usando la fórmula: sen(WCA) = (WV × sen(WA)) / TAS")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Donde: WA = Diferencia angular entre viento y rumbo = |270° - 220°| = 50°")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("sen(WCA) = (25 × sen 50°) / 120 = (25 × 0.7660) / 120 = 19.15 / 120 = 0.1596")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("WCA = arcsen(0.1596) ≈ 9.2° → WCA ≈ 9°")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

# Dirección del viento
p = doc.add_paragraph()
run = p.add_run("Dirección: El viento sopla desde 220° (SW), el rumbo es 270° (W). El viento tiene un componente desde la izquierda del rumbo. Por lo tanto, la corrección es LEFT (L).")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("WCA = 9° L (Left)")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("2. Cálculo de la Ground Speed (GS):")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("GS = TAS × cos(WCA) + WV × cos(WA)")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("GS = 120 × cos(9°) + 25 × cos(50°)")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("GS = 120 × 0.9877 + 25 × 0.6428 = 118.52 + 16.07 = 134.6 kts")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("GS ≈ 135 kts")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("3. Explicación del efecto del viento:")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("El viento de 220° con rumbo 270° genera dos componentes:")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("• Componente de cola (tailwind): 25 × cos(50°) ≈ 16 kts, que incrementa la GS por encima de la TAS (135 vs 120 kts).")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("• Componente cruzado (crosswind): 25 × sen(50°) ≈ 19 kts desde la izquierda, que requiere corrección de rumbo (WCA ≈ 9° L) para mantener el TC de 270°.")
run.font.size = Pt(10)
run.font.name = 'Arial'

doc.add_paragraph()

# P6
p = doc.add_paragraph()
run = p.add_run("P6. Gestión de Combustible (16 pts)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

p = doc.add_paragraph()
run = p.add_run("Datos: Consumo = 10 gal/h | Combustible útil = 56 gal | Búsqueda 90 min | Regreso 54 NM GS 135 kts | Reserva 45 min")
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Arial'

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("a) Combustible consumido durante la búsqueda (90 min):")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Procedimiento: Consumo por hora = 10 gal/h. Tiempo de búsqueda = 90 min = 1.5 h")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Gas_bsq = 10 gal/h × 1.5 h = 15 gal")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("b) Combustible para el regreso (54 NM a 135 kts GS):")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Procedimiento: Tiempo de regreso = Dist / GS = 54 / 135 = 0.4 h = 24 min")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Gas_reg = 10 gal/h × 0.4 h = 4 gal")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("c) Combustible de reserva legal (45 min):")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Procedimiento: Reserva = 45 min = 0.75 h")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Gas_res = 10 gal/h × 0.75 h = 7.5 gal")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("d) Total requerido vs Total a bordo:")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Gas_total = Gas_bsq + Gas_reg + Gas_res = 15 + 4 + 7.5 = 26.5 gal")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

p = doc.add_paragraph()
run = p.add_run("Combustible a bordo: 56 gal")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Combustible remanente: 56 - 26.5 = 29.5 gal")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("¿Suficiente? SÍ, hay excedente de 29.5 gal (más del 100% de reserva adicional sobre lo requerido).")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

doc.add_paragraph()

# P7
p = doc.add_paragraph()
run = p.add_run("P7. Corrección por Densidad Altitud (12 pts)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

p = doc.add_paragraph()
run = p.add_run("Datos: Altitud presión = 8,500 ft | OAT = +25 °C | Presión QNE (1013.25 hPa)")
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Arial'

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("a) Cálculo de la Densidad Altitud:")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Paso 1: Determinar la temperatura estándar ISA para 8,500 ft:")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("ISA STD a 8,500 ft = 15°C - (1.98°C × 8.5) = 15 - 16.83 = -1.83°C ≈ -2°C")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Paso 2: Calcular la desviación de temperatura (ISA Deviation):")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("ISA Dev = OAT - ISA STD = 25°C - (-2°C) = +27°C")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Paso 3: Aplicar la regla de 120 ft/°C (o usar el E6B):")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Density Altitude ≈ Pressure Altitude + (120 × ISA Dev)")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("DA ≈ 8,500 + (120 × 27) = 8,500 + 3,240 = 11,740 ft")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x00, 0x70, 0x30)

p = doc.add_paragraph()
run = p.add_run("(Usando el E6B, el resultado es aproximadamente 11,700 ft, consistente con el cálculo manual.)")
run.font.size = Pt(10)
run.font.name = 'Arial'
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("b) Pregunta de reflexión:")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("Efectos de despegar en un aeropuerto alto y caluroso (alta densidad altitud):")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("1. TAS real: AUMENTA. Con menor densidad del aire, la TAS para una misma IAS es mayor (se vuela más rápido respecto al aire, pero el desempeño aerodinámico se degrada).")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("2. Consumo de combustible: AUMENTA en términos de TAS, pero la eficiencia del motor disminuye. El motor recibe menos oxígeno, reduciendo su potencia. Para compensar, se requiere mayor tiempo de despegue y la tasa de ascenso es menor, lo que puede incrementar el consumo total en la fase de ascenso.")
run.font.size = Pt(10)
run.font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("3. Carrera de despegue: AUMENTA significativamente. La menor densidad del aire reduce la sustentación generada por las alas y la eficiencia de las hélices, requiriendo una mayor velocidad de despegue (TAS) y, por lo tanto, una pista más larga. En condiciones extremas (alta DA), puede exceder la longitud de pista disponible.")
run.font.size = Pt(10)
run.font.name = 'Arial'

doc.add_paragraph()

# ========== SUMMARY TABLE ==========
add_header_table(doc, "RESUMEN DE PUNTAJES", "")

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_headers = ["Sección", "Descripción", "Puntaje"]
for i, h in enumerate(summary_headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "D6E4F0")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'

summary = [
    ("A", "Cartas Aeronáuticas", "20"),
    ("B", "Sistemas y Radionavegación", "20"),
    ("C", "Velocidades y Altimetría", "20"),
    ("D — P5", "WCA y GS (E6B)", "12"),
    ("D — P6", "Gestión de Combustible", "16"),
    ("D — P7", "Corrección por Densidad Altitud", "12"),
]
for i, (sec, desc, pts) in enumerate(summary):
    row = table.rows[i + 1]
    for j, txt in enumerate([sec, desc, pts]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        if j == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Total row
row_total = table.add_row()
cells_total = row_total.cells
for j, txt in enumerate(["", "TOTAL", "100"]):
    p = cells_total[j].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    if j == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = "/home/andres/Documentos/ANTIGRAVITY-PROYECTOS/NAVEGACION-AEREA/Solucion_Examen_Unidad2_Navegacion.docx"
doc.save(output_path)
print(f"✅ Solución generada: {output_path}")
