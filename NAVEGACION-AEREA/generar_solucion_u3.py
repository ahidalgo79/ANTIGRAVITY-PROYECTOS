"""Genera la solución detallada del examen de la Unidad 3 (AIP/PIA).

Reutiliza el formato del solucionario de la Unidad 2 pero lee los reactivos
desde reactivos_u3.json y agrega la justificación de cada respuesta.
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
REACTIVOS = json.loads((BASE / "reactivos_u3.json").read_text(encoding="utf-8"))

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_table(doc, title, subtitle=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1B3A6B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(16)
    run.font.name = "Arial"
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.color.rgb = RGBColor(0xE8, 0xF0, 0xFF)
        r2.font.size = Pt(11)
        r2.font.name = "Arial"
    doc.add_paragraph()


def title_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    return p


def q_text(num, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}  ")
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return p


def answer_para(text, color="007030"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(*(int(color[i:i + 2], 16) for i in (0, 2, 4)))
    return p


def explanation_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.italic = True
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return p


# ========== TITLE ==========
add_header_table(
    doc,
    "SOLUCIÓN DETALLADA — EXAMEN UNIDAD 3",
    "Publicación de Información Aeronáutica (AIP/PIA) • Servicios de Información Aeronáutica",
)

p = doc.add_paragraph()
run = p.add_run("Instrucciones para el docente:")
run.bold = True
run.font.size = Pt(12)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

for item in [
    "Cada respuesta correcta tiene el puntaje indicado en el examen.",
    "Las justificaciones se basan en el Anexo 15 OACI y el Doc 8126 (Manual de servicios de información aeronáutica).",
    "Puntaje total: 100 puntos. Mínimo aprobatorio: 70 puntos.",
]:
    bullet(item, size=10)

# ========== SECTION A ==========
add_header_table(doc, "SECCIÓN A: Opción Múltiple — AIP/PIA", "40 puntos")

title_para("P1. Opción múltiple (40 pts — 2 pts c/u)")

# Justificaciones de cada pregunta de opción múltiple
JUST_OM = [
    "Tanto en español (PIA) como en inglés (AIP) se refieren al manual básico de información aeronáutica.",
    "El Anexo 15 establece que cada Estado publica su propia AIP a través de su Servicio de Información Aeronáutica (AIS). La OACI solo emite normas (SARPs).",
    "El Anexo 15 de la OACI (Servicios de Información Aeronáutica) es el que norma la producción y difusión de la AIP.",
    "La estructura de la AIP consta de tres partes: Parte 1 GEN (General), Parte 2 ENR (En-ruta) y Parte 3 AD (Aeródromos).",
    "La Parte ENR contiene rutas ATS, estructura del espacio aéreo, ayudas radioeléctricas y todos los datos necesarios para la fase en ruta.",
    "Los cambios temporales con duración superior a tres meses se publican en el Suplemento AIP. Los cambios de corta duración van en NOTAM.",
    "El NOTAM Clase I se distribuye inmediatamente a todas las estaciones y aeronaves; el Clase II solo se distribuye sobre solicitud.",
    "AIRAC establece fechas efectivas comunes con antelación suficiente para que todos los usuarios puedan preparar los cambios.",
    "Una Enmienda AIRAC introduce un cambio permanente en la AIP con una fecha efectiva coordinada (WEF = With Effect From).",
    "Las AIC se usan para información administrativa y de larga duración que no encaja en NOTAM, Suplemento o Enmienda de la AIP.",
    "El ciclo AIRAC estándar de OACI es de 28 días, lo que permite notificar los cambios con suficiente antelación.",
    "La Parte GEN contiene la legislación, reglamentos, acuerdos, unidades de medida y procedimientos generales del Estado.",
    "Las Enmiendas de la AIP se identifican por su impresión en papel azul.",
    "WEF (With Effect From) indica la fecha y hora desde la cual un cambio aeronáutico entra en vigor.",
    "El Suplemento AIP se publica en papel de color rosa (pink).",
    "NOTAMN (New) anuncia información nueva, como el establecimiento de una nueva ayuda; NOTAMR es de reemplazo y NOTAMC de cancelación.",
    "SNOWTAM es la serie especial de NOTAM que difunde condiciones de nieve, hielo y agua estancada en pistas.",
    "Los NOTAM los origina y distribuye el Servicio de Información Aeronáutica (AIS) del Estado.",
    "Las AIC (Circulares de Información Aeronáutica) se imprimen en papel verde.",
    "La AIP NO contiene pronósticos meteorológicos; esa información corresponde a los servicios meteorológicos (MET).",
]

for i, q in enumerate(REACTIVOS["opcion_multiple"]):
    q_text(f"{i + 1}.", q["texto"])
    for op in q["opciones"]:
        bullet(f"☐ {op}", size=10)
    correct = q["opciones"][q["correcta"]]
    answer_para(f"✅ Respuesta correcta: {correct}")
    explanation_para(f"Justificación: {JUST_OM[i]}")
    doc.add_paragraph()

# ========== SECTION B ==========
add_header_table(doc, "SECCIÓN B: Verdadero o Falso", "30 puntos")

title_para("P2. Verdadero o Falso (30 pts — 2 pts c/u)")

JUST_FV = [
    "Verdadero. La AIP es el manual básico de información aeronáutica indispensable para la navegación aérea y las operaciones aeroportuarias.",
    "Verdadero. La Parte AD reúne la información de aeródromos, helipuertos y sus procedimientos operacionales.",
    "Falso. Cada Estado es responsable de publicar su propia AIP; la OACI solo establece las normas (Anexo 15).",
    "Falso. Los NOTAM se difunden por vía electrónica (AFTN, internet) de forma inmediata; el papel no es su medio.",
    "Falso. El Suplemento AIP se usa para cambios TEMPORALES de larga duración (más de tres meses), no para cambios permanentes.",
    "Verdadero. La AIP debe estar siempre disponible y actualizada para todos los usuarios de la navegación aérea.",
    "Verdadero. Las Enmiendas AIRAC entran en vigor en fechas efectivas coordinadas internacionalmente (fechas AIRAC).",
    "Verdadero. El NOTAM, por ser más reciente y urgente, prevalece sobre la AIP hasta que esta se actualice formalmente.",
    "Verdadero. El Anexo 15 de la OACI regula los Servicios de Información Aeronáutica.",
    "Falso. La AIP es un documento público y de libre consulta para los usuarios de la navegación aérea.",
    "Verdadero. El ciclo AIRAC estándar se basa en fechas efectivas separadas 28 días.",
    "Verdadero. Las Enmiendas de la AIP se identifican por su papel de color azul.",
    "Falso. Los Suplementos AIP se imprimen en papel rosa (pink), no verde; el verde corresponde a las AIC.",
    "Verdadero. Las Circulares de Información Aeronáutica (AIC) se imprimen en papel verde.",
    "Falso. Los NOTAM se emiten principalmente para cambios temporales y urgentes, no para cambios permanentes.",
]

for i, q in enumerate(REACTIVOS["verdadero_falso"]):
    q_text(f"{i + 1}.", q["texto"])
    answer_para(f"  ( {'V' if q['correcta'] else 'F'} )")
    explanation_para(f"  {JUST_FV[i]}")
    doc.add_paragraph()

# ========== SECTION C ==========
add_header_table(doc, "SECCIÓN C: Caso de Estudio", "30 puntos")

title_para("📌  CASO PRÁCTICO: VERIFICACIÓN DE INFORMACIÓN AERONÁUTICA")

p = doc.add_paragraph()
run = p.add_run("Escenario:")
run.bold = True
run.font.size = Pt(11)
run.font.name = "Arial"
p = doc.add_paragraph()
run = p.add_run(REACTIVOS["caso_estudio"]["escenario"])
run.font.size = Pt(10)
run.font.name = "Arial"
doc.add_paragraph()

title_para("P3. Responder a partir del escenario (30 pts — 5 pts c/u)")

JUST_CASO = [
    "La información permanente de ruta y aeródromos se encuentra en la AIP y sus enmiendas; los NOTAM solo cubren cambios temporales y urgentes.",
    "Un cierre temporal de pista de dos semanas es un cambio temporal de corta duración que se difunde con un NOTAM de forma inmediata y prioritaria.",
    "El cambio permanente del procedimiento de aproximación se publica mediante una Enmienda AIRAC de la AIP con fecha efectiva coordinada.",
    "WEF (With Effect From) es la fecha efectiva AIRAC en la que el cambio entra en vigor; se publica con antelación para que todos los usuarios la conozcan.",
    "La información administrativa de larga duración (horarios, procedimientos administrativos) se publica en una Circular de Información Aeronáutica (AIC).",
    "La cancelación de un NOTAM se realiza con un NOTAMC (Cancellation), que anula la vigencia del NOTAM original.",
]

for i, q in enumerate(REACTIVOS["caso_estudio"]["preguntas"]):
    q_text(f"{i + 1}.", q["texto"])
    for op in q["opciones"]:
        bullet(f"☐ {op}", size=10)
    correct = q["opciones"][q["correcta"]]
    answer_para(f"✅ Respuesta correcta: {correct}")
    explanation_para(f"Justificación: {JUST_CASO[i]}")
    doc.add_paragraph()

# ========== SUMMARY TABLE ==========
add_header_table(doc, "RESUMEN DE PUNTAJES", "")

table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["Sección", "Descripción", "Puntaje"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "D6E4F0")
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

total_om = sum(q["puntos"] for q in REACTIVOS["opcion_multiple"])
total_fv = sum(q["puntos"] for q in REACTIVOS["verdadero_falso"])
total_caso = sum(q["puntos"] for q in REACTIVOS["caso_estudio"]["preguntas"])

summary = [
    ("A", "Opción Múltiple (20 reactivos)", str(total_om)),
    ("B", "Verdadero o Falso (15 reactivos)", str(total_fv)),
    ("C", "Caso de Estudio (6 reactivos)", str(total_caso)),
]
for i, (sec, desc, pts) in enumerate(summary):
    row = table.rows[i + 1]
    for j, txt in enumerate([sec, desc, pts]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        if j == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

row_total = table.add_row()
for j, txt in enumerate(["", "TOTAL", str(total_om + total_fv + total_caso)]):
    p = row_total.cells[j].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    if j == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = BASE / "Solucion_Examen_Unidad3_Navegacion.docx"
doc.save(output_path)
print(f"Solucionario generado exitosamente en {output_path.name}")
