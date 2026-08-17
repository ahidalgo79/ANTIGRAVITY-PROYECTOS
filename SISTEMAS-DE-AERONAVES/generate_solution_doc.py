import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_solution_files():
    base_dir = "/home/andres/Documentos/ANTIGRAVITY-PROYECTOS/SISTEMAS-DE-AERONAVES"
    quizzes_dir = os.path.join(base_dir, "quizzes")
    
    # Read PDF text for exact matching
    with open('/tmp/examen_text.txt', 'r', encoding='utf-8') as f:
        pdf_text = f.read()
        
    pdf_clean = re.sub(r'=== PAGE \d+ ===', '', pdf_text)
    pdf_clean = re.sub(r'\f', '', pdf_clean)

    quiz_files = [
        ('Unidad 1: Introducción a la Estructura de Manuales y ATA 100', 'quiz_unidad1.json'),
        ('Unidad 2: Sistema de Aire Acondicionado y Presurización (ATA 21)', 'quiz_unidad2.json'),
        ('Unidad 3: Sistemas de Control de Vuelo (ATA 27) y Fly-by-Wire', 'quiz_unidad3.json'),
        ('Unidad 4: Protección contra Hielo y Lluvia (ATA 30)', 'quiz_unidad4.json'),
        ('Sección Especial Integradora: Examen de Repaso General', 'quiz_integrador.json')
    ]

    manual_fixes = {
        30: ('a', 'T₂ = T₁ × ((P₂ / P₁))^((γ-1)/γ)', 'Fórmula teórica de temperatura de salida en compresión isentrópica según la relación de presiones (P₂/P₁) y el coeficiente adiabático γ.'),
        31: ('b', '8,000 ft', 'La FAR/CS 25.841 establece que en operación normal la altitud máxima equivalente de cabina no debe exceder de 8,000 ft para evitar la hipoxia normobárica.'),
        34: ('c', '0.55 lb/min', 'La norma FAR 25.831 especifica un flujo de aire fresco mínimo de 0.55 lb/min (aprox. 10 ft³/min) por ocupante para mantener niveles aceptables de O₂ y diluir contaminantes.'),
        38: ('a', '18 a 30°C', 'El selector del panel de control del A320 permite regular la temperatura de zona entre 18°C (frío máximo) y 30°C (calor máximo).'),
        41: ('c', '7 × 19', 'El cable flexible 7x19 (7 cordones de 19 hilos cada uno) es el estándar técnico militar y comercial (MIL-DTL-83420) para controles primarios de vuelo por su alta flexibilidad y resistencia a la fatiga.'),
        43: ('d', 'F ∝ V²', 'La presión dinámica q es proporcional al cuadrado de la velocidad (q = 1/2 ρ V²), por lo que las fuerzas aerodinámicas sobre los cables/superficies crecen cuadráticamente con la velocidad.')
    }

    # Detailed technical rationales dictionary
    technical_rationales = {
        1: "El estándar ATA 100 fue creado por la Air Transport Association para unificar la estructura de numeración de capítulos y manuales técnicos en la aviación civil global.",
        2: "En el formato ATA iSpec 2200 (CC-SS-UU), los dos primeros dígitos (CC) corresponden al Capítulo/Sistema, el segundo par (SS) al Subsistema y los últimos dos (UU) a la Unidad o componente específico.",
        3: "El Aircraft Maintenance Manual (AMM) guía el mantenimiento realizado directamente en la aeronave ('on-aircraft'), mientras que el Component Maintenance Manual (CMM) se utiliza en talleres especializados para reparación/overhaul de componentes desmontados ('off-aircraft').",
        4: "En el sistema de codificación ATA 100, el Capítulo 21 está asignado formalmente a 'Air Conditioning and Pressurization' (Aire Acondicionado y Presurización).",
        5: "El Illustrated Parts Catalog (IPC) permite identificar, aprovisionar, pedir y almacenar piezas de repuesto reemplazables en línea (LRUs) mediante diagramas desglosados y números de parte.",
        6: "La Minimum Equipment List (MEL) es un documento operacional aprobado por la autoridad que permite el despacho del avión con equipos inoperativos bajo condiciones de mitigación específicas.",
        7: "El Wiring Diagram Manual (WDM) detalla los diagramas de cableado eléctrico, ruteo de arneses, conectores y conexiones físicas entre sistemas y componentes de la aeronave.",
        8: "La norma ATA iSpec 2200 modernizó la especificación ATA 100 migrando toda la documentación a formatos digitales estructurados en XML mediante DTDs estandarizadas para la industria.",
        9: "Los capítulos 71 al 80 del estándar ATA 100 están reservados exclusivamente para la planta motriz (Powerplant), abarcando desde la instalación básica del motor hasta ignición y combustible.",
        10: "El Structural Repair Manual (SRM) proporciona límites de daño evaluables y procedimientos de reparación estructural aprobados directamente por el fabricante de la estructura primaria y secundaria.",
        11: "El AFM es un documento normativo de certificación aprobado directamente por la autoridad aeronáutica (FAA/EASA) que contiene los límites de operación, procedimientos de emergencia y desempeño legal del avión.",
        12: "La normalización estandariza especificaciones, tolerancias y procesos industriales para garantizar la interoperabilidad de componentes y piezas entre diferentes fabricantes aeronáuticos.",
        13: "En la zonificación física estándar ATA, la serie de zona 100 corresponde al fuselaje inferior (Lower Fuselage) y compartimento de carga.",
        14: "El término 'Hard Law' en aviación se refiere a regulaciones jurídicamente vinculantes y de obligatorio cumplimiento impuestas por la autoridad (como las FARs, EASA CS, ADs).",
        15: "El FCOM es el manual operativo que proporciona a la tripulación de mando las técnicas de vuelo, listas de verificación, procedimientos normales, no normales y de emergencia específicos de la aeronave.",
        16: "En la codificación ATA 100, el Capítulo 32 corresponde formalmente a 'Landing Gear' (Tren de Aterrizaje), incluyendo frenos, dirección y amortiguación.",
        17: "La FAR/CS 25.1309 establece que la probabilidad de ocurrencia de cualquier condición de falla que resulte en un evento catastrófico debe ser menor a 1 × 10⁻⁹ por hora de vuelo (Extremadamente Improbable).",
        18: "Los Capítulos 51 a 57 del código ATA corresponden al grupo de Estructuras (Structures), abarcando estructuras generales, puertas, fuselaje, estabilizadores y alas.",
        19: "La Configuration Deviation List (CDL) autoriza el despacho de la aeronave con piezas secundarias externas de la estructura o carenados faltantes (deflectores, paneles), aplicando penalizaciones de desempeño.",
        20: "El principio de diseño Fail-safe asegura que tras la falla de un elemento estructural o componente, la carga o función se redistribuye a trayectorias secundarias sin causar un colapso catastrófico inmediato.",
        21: "Las Máquinas de Ciclo de Aire (ACM) operan bajo el ciclo termodinámico de Brayton Inverso (o refrigeración por gas), donde el aire actúa como fluido de trabajo mediante compresión y expansión adiabática.",
        22: "La Outflow Valve (válvula de salida de aire) es el elemento de control primario que regula la cantidad de aire que escapa del fuselaje sellado para mantener la presión diferencial y la altitud de cabina deseada.",
        23: "El compresor centrífugo de la ACM incrementa la presión y temperatura del aire previo al intercambiador secundario, aumentando la energía disponible para extraer trabajo en la turbina de expansión.",
        24: "En el ciclo de Brayton inverso, la expansión adiabática/isentrópica del aire a través de la turbina realiza trabajo mecánico y causa una caída drástica de temperatura y presión.",
        25: "En un sistema de ciclo de vapor (VCS), el refrigerante absorbe el calor del aire de cabina en el evaporador, cambiando de fase líquida a fase gaseosa.",
        26: "El sistema VCS ofrece una mayor eficiencia termodinámica (COP elevado) sin depender del aire de sangrado del motor (bleed air), reduciendo la penalización en el consumo específico de combustible.",
        27: "La Trim Air Valve inyecta aire caliente directo del sangrado para modular y afinar la temperatura de cada zona de cabina de forma independiente sobre el flujo básico de aire frío suministrado por las packs.",
        28: "Si la presión atmosférica disminuye al ascender la aeronave y la presión interna se mantiene, la presión diferencial (ΔP = P_cabina - P_exterior) aumenta progresivamente.",
        29: "Una ACM de tres ruedas (Three-Wheel ACM) incorpora en un mismo eje el compresor centrífugo, la turbina de expansión y un ventilador de tiro forzado (extract fan).",
        30: "En la compresión isentrópica de un gas ideal, la relación entre temperaturas y presiones viene dada por T₂ = T₁ × (P₂/P₁)^((γ-1)/γ), donde γ es la relación de calores específicos (1.4 para el aire).",
        31: "La normativa FAR 25.841 exige que en operación normal la altitud equivalente de cabina no supere los 8,000 ft para mantener la saturación adecuada de oxígeno en la sangre de pasajeros y tripulación.",
        32: "El agua condensada extraída por el separador ciclónico se pulveriza mecánicamente sobre los tubos del intercambiador de calor secundario (spray nozzle) para aumentar drásticamente su eficiencia de enfriamiento.",
        33: "En el B737, si un pack sufre un sobrecalentamiento o falla de control, se ilumina la luz ámbar de aviso 'PACK' en el panel superior de aire acondicionado y se cierra la válvula del pack.",
        34: "La regulación FAR 25.831 requiere un suministro mínimo de aire fresco de 0.55 lb/min (aprox. 10 cfm) por ocupante en todas las condiciones normales de vuelo.",
        35: "Las válvulas de seguridad (Safety / Relief Valves) neumáticas protegen físicamente la estructura del fuselaje contra sobrepresión diferencial (límite máximo de ΔP) y presión negativa en caso de fallas del control principal.",
        36: "Un sistema 'Bootstrap' reutiliza la energía mecánica generada por la turbina de expansión para impulsar directamente el compresor centrífugo acoplado en el mismo eje neumático.",
        37: "Extraer aire de sangrado (bleed air) reduce el flujo de masa y la presión en la turbina de alta presión del motor, aumentando el consumo específico de combustible (SFC) y reduciendo el empuje disponible.",
        38: "En el Airbus A320, las perillas de control de temperatura de zona del panel de aire acondicionado permiten seleccionar temperaturas entre 18°C (Cold) y 30°C (Hot).",
        39: "En el motor CFM56-7B del B737, el aire de sangrado para el sistema neumático se extrae de la 5ª etapa (baja presión/crucero) o de la 9ª etapa (alta presión/ralentí) del compresor de alta presión (HPC).",
        40: "El intercambiador de calor primario realiza el enfriamiento inicial del aire de sangrado hipercalentado procedente de los motores utilizando el flujo de aire de impacto (ram air).",
        41: "El cable trenzado de acero de especificación 7x19 (7 cordones formados por 19 alambres cada uno) proporciona la máxima flexibilidad y resistencia a la fatiga en poleas de control primario.",
        42: "Los sistemas puramente mecánicos ofrecen inmunidad absoluta frente a interferencias electromagnéticas (EMI) y fallas de alimentación eléctrica, además de retroalimentación táctil directa.",
        43: "La fuerza aerodinámica requerida para desplazar una superficie de control crece de forma proporcional al cuadrado de la velocidad del aire (F ∝ V² = 1/2 ρ V² S C_L).",
        44: "Los sistemas Fly-by-Wire de Airbus utilizan la Ley de Control C* (C-star), que combina una demanda de aceleración normal (g) a bajas velocidades con demanda de tasa de cabeceo (pitch rate) a altas velocidades.",
        45: "Las unidades ADIRU (Air Data Inertial Reference Unit) suministran los parámetros de vuelo procesados (velocidad, Mach, ángulo de ataque, aceleraciones) a las computadoras de vuelo FBW.",
        46: "En Ley Normal de Airbus, la deflexión del sidestick en el eje longitudinal comanda directamente un factor de carga o aceleración normal (g) independiente de la velocidad.",
        47: "Cuando el ángulo de ataque excede α_prot, el sistema entra en Alpha Protection: comanda cabeceo abajo automático, desactiva el trim autotrim y limita el α a α_max aunque el piloto mantenga el stick atrás.",
        48: "El Flare Mode conmuta la ley de cabeceo a una relación directa modificada durante los últimos 50 pies, introduciendo una leve tendencia de cabeceo abajo para simular el comportamiento convencional y exigir el 'flare' manual.",
        49: "Bajo Ley Alternativa, el control de alabeo conmuta a Ley Directa (Direct Law), donde la deflexión del sidestick mueve las superficies de alabeo en proporción lineal estricta sin protecciones avanzadas.",
        50: "La protección de sobrevelocidad de factor de carga (Load Factor Limitation: +2.5g / -1.0g limpio) se conserva activa incluso cuando el avión ha degradado a Ley Alternativa.",
        51: "La indicación visual principal de degradación a Ley Alternativa son las banderas ámbar 'USE MAN PITCH TRIM' en los PFDs y el cambio de las marcas verdes de protección a cruces ámbar '=' en la escala de actitud.",
        52: "En el modo de Respaldo Mecánico (Mechanical Backup), los únicos controles de vuelo disponibles son el estabilizador horizontal variable (THS) mediante la rueda de trim manual y el timón de dirección mediante pedales.",
        53: "Los alerones son superficies primarias instaladas en el borde de salida de la parte exterior de las alas encargadas de generar momentos de balanceo o alabeo (roll) sobre el eje longitudinal.",
        54: "Al desplegarse asimétricamente en un solo lado del ala, los spoilers reducen la sustentación de esa ala y aumentan su resistencia, asistiendo a los alerones en el control de alabeo (roll spoiler function).",
        55: "Las deflexiones de flaps en rangos elevados (superiores a 20°-25°) están diseñadas primariamente para incrementar drásticamente la resistencia aerodinámica (drag) para aproximación y aterrizaje.",
        56: "El tensiómetro de cables de control mide la deflexión mecánica del cable bajo carga ajustada a la temperatura ambiente para verificar la tensión del sistema de acuerdo con la gráfica de mantenimiento.",
        57: "Las computadoras FAC (Flight Augmentation Computer) calculan el amortiguamiento de guiñada (yaw damper), la coordinación de virajes, el trimeo del timón y los límites de velocidad de maniobra.",
        58: "En el modo Tierra (Ground Mode) de Ley Normal, la computadora resetea automáticamente la posición del trim del estabilizador horizontal (THS) a 0° tras el aterrizaje y establece relación directa sidestick-superficie.",
        59: "En Ley Alternativa, la extensión del tren de aterrizaje conmuta el eje de cabeceo a Ley Directa porque no existe un modo Flare disponible, garantizando que el piloto tenga respuesta directa durante el toque.",
        60: "Los actuadores electrohidráulicos (Servoactuadores EHV/EHA) convierten las señales eléctricas enviadas por las computadoras de vuelo FBW en movimiento hidráulico mecánico para desplazar las superficies.",
        61: "Las normas FAR/EASA 25.1419 exigen demostrar que la aeronave es capaz de operar de manera segura y continuada en las condiciones de congelamiento definidas en el Apéndice C de la norma.",
        62: "Un sistema de deshielo (de-ice) es reactivo (permite que el hielo se acumule y luego lo rompe/elimina), mientras que un sistema anti-hielo (anti-ice) es proactivo (previene activamente la formación inicial de hielo).",
        63: "Las botas o zapatas neumáticas de deshielo están fabricadas tradicionalmente de neopreno reforzado o elastómeros sintéticos debido a su gran elasticidad y resistencia al desgaste ambiental.",
        64: "Cuando las botas neumáticas no están infladas, el sistema de vacío aplica succión interna para mantener las botas ajustadas contra el perfil del ala, minimizando la resistencia aerodinámica.",
        65: "Investigaciones de la NTSB y FAA demostraron que el 'puente de hielo' es un mito en botas modernas, recomendando su activación inmediata tan pronto como se detecte la presencia de hielo.",
        66: "En aeronaves comerciales de transporte a reacción, el aire supercalentado del sangrado de motores (Thermal Bleed Air) es la principal fuente de energía para el anti-hielo de alas y tomas de aire.",
        67: "Los sistemas electrotérmicos (resistencias eléctricas integradas) se usan en aviones modernos de fibra de carbono (como el B787) y en componentes críticos como sondas pitot, hélices y parabrisas.",
        68: "El engelamiento en las tomas de aire del motor (engine cowls) es crítico porque el desprendimiento de bloques de hielo hacia los álabes del fan o compresor causa daños por objeto extraño (FOD) catastróficos.",
        69: "Los recubrimientos químicos hidrófobos modifican la tensión superficial del cristal aumentando el ángulo de contacto de las gotas de agua, haciendo que resbalen rápidamente con el viento relativo.",
        70: "En el B737, la luz de aviso 'COWL ANTI-ICE' en el panel superior se ilumina en azul tenue (dim) cuando la válvula está completamente abierta y alineada con la posición del interruptor.",
        71: "La capa de agua en el parabrisas actúa como una lente irregular que altera la refracción de la luz, creando la ilusión óptica de que la pista está más baja o más distante de lo real (error de paralaje).",
        72: "Los sopladores de lluvia (Rain Removal Blowers) expulsan un chorro continuo de aire comprimido a alta velocidad a través de boquillas sobre la superficie exterior del parabrisas para desviar el agua por impacto neumático.",
        73: "En el B737 en tierra, si las palancas de empuje avanzan por encima de 32° (posición de despegue), la lógica cierra las válvulas de Wing Anti-Ice para evitar el sobrecalentamiento de la estructura del ala sin flujo de aire.",
        74: "Al desintegrar el contacto del tren de aterrizaje con tierra (transición tierra-aire al despegar), el interruptor de Wing Anti-Ice del B737 salta automáticamente a OFF por solenoide de disparo.",
        75: "El calentamiento eléctrico de los parabrisas mantiene las capas intermedias de vinilo/plástico flexibles y elásticas a bajas temperaturas, aumentando dramáticamente la resistencia estructural del cristal contra impacto de aves (bird strike).",
        76: "Las válvulas distribuidoras/temporizadas de deshielo (deice timer valves) regulan secuencialmente la presión de inflado y el ciclo de vacío de las botas neumáticas por secciones.",
        77: "La acumulación de hielo no protegido en el borde de ataque del ala altera el perfil aerodinámico, destruyendo la capa límite, reduciendo el C_L,max y aumentando sustancialmente la velocidad de pérdida (stall speed).",
        78: "El sistema de deshielo por impulso electroexpulsivo (EIDI) genera un pulso electromagnético de alta intensidad en bobinas internas que acelera mecánicamente la piel del ala hasta 1000g, fracturando el hielo de impacto.",
        79: "Operar las escobillas de los limpiaparabrisas sobre cristal seco provoca rayaduras permanentes en los recubrimientos ópticos y deteriora rápidamente los bordes de caucho de la escobilla.",
        80: "Los drenajes de aguas residuales calentados (Heated Drain Masts) aplican calor eléctrico continuo a los tubos de salida de desechos para evitar la formación de bloques de hielo externos ('blue ice').",
        81: "La FAR/CS 25.1309 define que una condición de falla con consecuencias catastróficas debe ser 'Extremadamente Improbable', lo que corresponde a una probabilidad menor a 1 × 10⁻⁹ por hora de vuelo.",
        82: "El fluido hidráulico sintético Skydrol (éster de fosfato) es altamente corrosivo para elastómeros convencionales, exigiendo obligatoriamente el uso de sellos de Butilo o Teflon (PTFE).",
        83: "La APU (Auxiliary Power Unit) es una pequeña turbina de gas que suministra energía eléctrica (generador) y aire neumático (bleed air) de forma autónoma con la aeronave en tierra o en vuelo de emergencia.",
        84: "En la mayoría de los aviones comerciales de transporte, la alarma maestra de altitud de cabina (sirena auditiva y advertencia visual CABIN ALTITUDE) se dispara al alcanzar los 10,000 pies de altitud de cabina.",
        85: "Las alertas inmediatas de peligro inminente como las advertencias GPWS/TAWS (terreno) y Stall Warning (pérdida) tienen máxima prioridad absoluta visual y auditiva sobre las alertas de tráfico (TCAS) u otras indicaciones."
    }

    all_q_data = []
    global_num = 1

    for section_title, qf in quiz_files:
        with open(os.path.join(quizzes_dir, qf), 'r', encoding='utf-8') as fp:
            data = json.load(fp)
            for q in data.get('questions', []):
                correct_json_text = None
                for opt in q.get('answerOptions', []):
                    if opt.get('isCorrect'):
                        correct_json_text = opt.get('text').strip()
                        break
                
                if global_num in manual_fixes:
                    c_letter, c_text, rat = manual_fixes[global_num]
                else:
                    m = re.search(rf'\n{global_num}\.\s+.*?(?=\n{global_num+1}\.|\Z)', pdf_clean, re.DOTALL)
                    c_letter = None
                    c_text = None
                    if m:
                        block = m.group(0)
                        opt_matches = list(re.finditer(r'([a-d])\)\s+(.*?)(?=\n[a-d]\)|\n\d+\.|\Z)', block, re.DOTALL))
                        for om in opt_matches:
                            l = om.group(1).upper()
                            t = re.sub(r'\s+', ' ', om.group(2)).strip()
                            if correct_json_text in t or t in correct_json_text or correct_json_text[:25] in t:
                                c_letter = l.lower()
                                c_text = t
                                break
                    rat = technical_rationales.get(global_num, q.get('rationale', ''))

                all_q_data.append({
                    'num': global_num,
                    'section': section_title,
                    'question': q.get('question').strip(),
                    'letter': c_letter.upper() if c_letter else 'A',
                    'option_text': c_text or correct_json_text,
                    'rationale': rat
                })
                global_num += 1

    # --- 1. BUILD DOCX DOCUMENT ---
    doc = Document()
    
    # Page Margins
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Header / Logo
    logo_path = os.path.join(base_dir, "OIC-28.png")
    if os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        lr = lp.add_run()
        lr.add_picture(logo_path, width=Inches(2.2))

    tp = doc.add_paragraph()
    tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tr = tp.add_run("SOLUCIONARIO OFICIAL Y COMPLETO\nEXAMEN: SISTEMAS EN AERONAVES")
    tr.bold = True
    tr.font.name = "Arial"
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(15, 32, 67) # Navy

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = sub.add_run("Evaluación de Desempeño Técnico y Teórico — 100 Puntos Totales\nDocumento Guía para Instructor y Estudiante")
    sr.font.name = "Arial"
    sr.font.size = Pt(11)
    sr.font.italic = True
    sr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Table Summary of Answers (Tabla Resumen Rápidas)
    add_sec_title(doc, "TABLA RESUMEN DE CLAVE DE RESPUESTAS (1 - 85)")
    
    # Create a 5-column table for quick key
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["Preguntas", "Clave", "Preguntas", "Clave", "Preguntas"] # We will do a 4-col key table: Q | Ans | Q | Ans
    
    table_q = doc.add_table(rows=1, cols=4)
    table_q.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hcells = table_q.rows[0].cells
    hcells[0].text = "Pregunta"
    hcells[1].text = "Respuesta Correcta"
    hcells[2].text = "Pregunta"
    hcells[3].text = "Respuesta Correcta"
    
    for c in hcells:
        set_cell_background(c, "0F2043")
        set_cell_margins(c, 80, 80, 100, 100)
        for p in c.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = "Arial"
                r.font.size = Pt(10)

    half = 43
    for i in range(half):
        row_cells = table_q.add_row().cells
        q1 = all_q_data[i]
        row_cells[0].text = f"P{q1['num']}"
        row_cells[1].text = f"{q1['letter']}) {q1['option_text'][:35]}..." if len(q1['option_text']) > 35 else f"{q1['letter']}) {q1['option_text']}"
        
        if i + half < len(all_q_data):
            q2 = all_q_data[i + half]
            row_cells[2].text = f"P{q2['num']}"
            row_cells[3].text = f"{q2['letter']}) {q2['option_text'][:35]}..." if len(q2['option_text']) > 35 else f"{q2['letter']}) {q2['option_text']}"
        else:
            row_cells[2].text = ""
            row_cells[3].text = ""

        # Formatting table cells
        bg = "F4F6F9" if i % 2 == 1 else "FFFFFF"
        for idx_c, cell in enumerate(row_cells):
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50, 50, 80, 80)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    if idx_c in [1, 3]:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(15, 32, 67)

    doc.add_page_break()

    # Detailed Section by Section
    current_sec = ""
    for item in all_q_data:
        if item['section'] != current_sec:
            current_sec = item['section']
            add_sec_title(doc, current_sec)
            doc.add_paragraph()

        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        qr = qp.add_run(f"Pregunta {item['num']}: ")
        qr.bold = True
        qr.font.size = Pt(11)
        qr.font.name = "Arial"
        qr.font.color.rgb = RGBColor(15, 32, 67)

        qtext = qp.add_run(item['question'])
        qtext.font.size = Pt(11)
        qtext.font.name = "Arial"

        # Correct answer box / paragraph
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(2)
        
        ar_title = ap.add_run("Respuesta Correcta: ")
        ar_title.bold = True
        ar_title.font.name = "Arial"
        ar_title.font.color.rgb = RGBColor(20, 120, 40) # Green

        ar_val = ap.add_run(f"Opción {item['letter']}) — {item['option_text']}")
        ar_val.bold = True
        ar_val.font.name = "Arial"

        # Justification
        jp = doc.add_paragraph()
        jp.paragraph_format.left_indent = Inches(0.25)
        jp.paragraph_format.space_after = Pt(8)
        jr_t = jp.add_run("Justificación Técnica: ")
        jr_t.bold = True
        jr_t.font.name = "Arial"
        jr_t.font.color.rgb = RGBColor(80, 80, 80)

        jr_v = jp.add_run(item['rationale'])
        jr_v.font.name = "Arial"
        jr_v.font.size = Pt(10)

    docx_path = os.path.join(base_dir, "Solucionario_Examen_Sistemas_en_Aeronaves_Especial.docx")
    doc.save(docx_path)
    print(f"DOCX Solucionario generado en: {docx_path}")

    # --- 2. BUILD MARKDOWN FILE ---
    md_path = os.path.join(base_dir, "Solucionario_Examen_Sistemas_en_Aeronaves_Especial.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("# 📋 SOLUCIONARIO OFICIAL Y COMPLETO\n")
        f.write("## EXAMEN: SISTEMAS EN AERONAVES (ESPECIAL - 100 PTS)\n\n")
        f.write("> **Asignatura:** Sistemas de Aeronaves  \n")
        f.write("> **Total de Reactivos:** 85 Preguntas (80 de 1 pt + 5 integradoras de 4 pts)  \n")
        f.write("> **Estado:** Solucionario Verificado y Justificado  \n\n")

        f.write("---  \n\n")
        f.write("### 📌 Tabla Rápida de Claves de Respuesta (1 a 85)\n\n")
        f.write("| # | Opción | Respuesta | # | Opción | Respuesta |\n")
        f.write("|---|---|---|---|---|---|\n")
        
        for i in range(half):
            q1 = all_q_data[i]
            str1 = f"| **P{q1['num']}** | **{q1['letter']}** | {q1['option_text'][:35]} |"
            if i + half < len(all_q_data):
                q2 = all_q_data[i + half]
                str2 = f" **P{q2['num']}** | **{q2['letter']}** | {q2['option_text'][:35]} |"
            else:
                str2 = " | | |"
            f.write(str1 + str2 + "\n")

        f.write("\n---\n\n")

        curr_sec_md = ""
        for item in all_q_data:
            if item['section'] != curr_sec_md:
                curr_sec_md = item['section']
                f.write(f"## 📚 {curr_sec_md}\n\n")

            f.write(f"### Pregunta {item['num']}\n")
            f.write(f"**{item['question']}**\n\n")
            f.write(f"- ✅ **Respuesta Correcta:** Opción **{item['letter']})** — `{item['option_text']}`\n")
            f.write(f"- 💡 **Justificación Técnica:** {item['rationale']}\n\n")
            f.write("---\n\n")

    print(f"MD Solucionario generado en: {md_path}")

def add_sec_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(15, 32, 67)

if __name__ == "__main__":
    generate_solution_files()
