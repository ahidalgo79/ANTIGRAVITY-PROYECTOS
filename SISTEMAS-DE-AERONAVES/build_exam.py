import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def clear_document(doc):
    """Remove all paragraphs and tables from the document body."""
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    for table in doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)
        tbl._tbl = tbl._element = None

def get_quizzes(quizzes_dir):
    """Load and sort the quizzes."""
    quiz_files = [
        "quiz_unidad1.json",
        "quiz_unidad2.json",
        "quiz_unidad3.json",
        "quiz_unidad4.json",
        "quiz_integrador.json"
    ]
    quizzes = []
    for q_file in quiz_files:
        path = os.path.join(quizzes_dir, q_file)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                quizzes.append({
                    "filename": q_file,
                    "title": data.get("title", ""),
                    "questions": data.get("questions", [])
                })
    return quizzes

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    # Basic formatting to match typical exam styles
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(14)
            
def add_question(doc, q_idx, q_data):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    run_num = p.add_run(f"{q_idx}. ")
    run_num.bold = True
    p.add_run(q_data["question"])
    
    options = ["a)", "b)", "c)", "d)", "e)", "f)"]
    
    for i, opt in enumerate(q_data.get("answerOptions", [])):
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent = Inches(0.5)
        p_opt.paragraph_format.space_after = Pt(2)
        p_opt.add_run(f"{options[i]} {opt['text']}")

def main():
    base_dir = "/home/andres/Documentos/ANTIGRAVITY-PROYECTOS/SISTEMAS-DE-AERONAVES"
    template_path = os.path.join(base_dir, "Plantilla_Examen.docx")
    quizzes_dir = os.path.join(base_dir, "quizzes")
    output_path = os.path.join(base_dir, "Examen_Sistemas_en_Aeronaves.docx")
    
    doc = Document(template_path)
    clear_document(doc)
    
    # Logo
    logo_path = os.path.join(base_dir, "OIC-28.png")
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logo_run = logo_p.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.0))
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run("EXAMEN: SISTEMAS EN AERONAVES")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Arial"
    
    doc.add_paragraph("Nombre del alumno: ____________________________________________________  Fecha: ____________")
    doc.add_paragraph()
    
    quizzes = get_quizzes(quizzes_dir)
    
    global_q_idx = 1
    answer_key = []
    
    for i, quiz in enumerate(quizzes):
        if "Integrador" in quiz['title']:
            add_heading(doc, f"SECCIÓN ESPECIAL: {quiz['title']}", level=2)
        else:
            unit_num = i + 1
            add_heading(doc, f"Unidad {unit_num}: {quiz['title']}", level=2)
        doc.add_paragraph()
        
        for q in quiz["questions"]:
            add_question(doc, global_q_idx, q)
            doc.add_paragraph() # spacing
            
            # Record for answer key
            options = ["A", "B", "C", "D", "E", "F"]
            correct_opt = ""
            rationale = ""
            for j, opt in enumerate(q.get("answerOptions", [])):
                if opt.get("isCorrect"):
                    correct_opt = options[j]
                    rationale = opt.get("rationale", "")
                    break
                    
            answer_key.append({
                "num": global_q_idx,
                "correct": correct_opt,
                "rationale": rationale,
                "question": q["question"]
            })
            global_q_idx += 1
            
    # Add Answer Key
    doc.add_page_break()
    key_title = doc.add_paragraph()
    key_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    kr = key_title.add_run("CLAVE DE RESPUESTAS (SOLO PARA EL PROFESOR)")
    kr.bold = True
    kr.font.size = Pt(16)
    
    for ans in answer_key:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_bold = p.add_run(f"Pregunta {ans['num']}: Respuesta {ans['correct']}. ")
        run_bold.bold = True
        p.add_run(f"Justificación: {ans['rationale']}")
        
    doc.save(output_path)
    print(f"Examen generado exitosamente en: {output_path}")

if __name__ == "__main__":
    main()
